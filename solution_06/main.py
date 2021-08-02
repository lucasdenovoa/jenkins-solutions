# coding=utf-8
import sys
import argparse
import os
import yaml
import shutil
import time
from docx import Document

dir_path = os.path.dirname(os.path.realpath(__file__))

parser = argparse.ArgumentParser(description="Run the script")

parser.add_argument(
            "-c",
            "--config",
            type=str,
            help="run the configuration script",
)

parser.add_argument(
            "-o",
            "--output",
            type=str,
            help="run ",
)
    
args = parser.parse_args()


with open(dir_path+ "/" + args.config, 'r') as file:
        bands_object = yaml.safe_load(file)

TOTAL_OF_BANDS = len(bands_object['bands'])
bands_name = []
bands_about = []
bands_member = []
bands_musics_name = []
bands_album = []

# Os dados das bandas sao armazenado aqui
for index_name in range(0, TOTAL_OF_BANDS):
    bands_name.append(bands_object['bands'][index_name]['name'])
    bands_about.append(bands_object['bands'][index_name]['about'])
    bands_musics_name.append(bands_object['bands'][index_name]['musics'])

#print(bands_musics_name[troca a banda][troca album])

if os.path.exists(dir_path + "/" + args.output):
    print "Esse diretório ja existe"
    
else:
    os.mkdir(dir_path + "/" + args.output)




#cria os arquivos no diretorio
for i in range(0, len(bands_name)):
    document = Document()
    name = bands_name[i].replace(' ', '')
    name = name.replace("'", '')
    print("Criando arquivos " + './'+ args.output +'/'+ name+".docx")


    document.add_heading('Bands '+ bands_name[i], level=1)
    document.add_paragraph(bands_about[i])  

    #Criando a tabela
    table = document.add_table(rows=len(bands_musics_name[0]), cols=2)

    hdr_cells = table.rows[0].cells

    document.add_heading('Musics '+ bands_name[i], level=2)
    #Mudar para tabela
    for x in range(0, 3):
        document.add_paragraph(bands_musics_name[i][x]['name'] + " " + bands_musics_name[i][x]['album'])  

    document.save('./'+ args.output +'/'+ name+ ".docx")
    

   


